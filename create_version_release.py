#!/usr/bin/env python3
import subprocess
import sys

# 确保安装了 python-docx
try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--user"])
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 标题
title = doc.add_paragraph("产品版本发布说明", style="Heading 1")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 102, 204)

# 各字段
fields = [
    ("产品名称", "明确是什么产品进行版本发布。"),
    ("版本号", "本次发布的具体版本编号。"),
    ("新功能", "详细描述新增的功能特性，例如操作流程、使用场景等。"),
    ("优化改进", "说明对现有功能在性能、界面、用户体验等方面做了哪些优化。"),
    ("修复的问题", "列举修复的各类漏洞、故障或错误。"),
    ("适用范围", "该版本适用的平台（如 Windows、iOS 等）、系统环境等。"),
    ("注意事项", "如是否存在已知问题、更新限制条件等。")
]

for field_name, field_desc in fields:
    # 字段名作为 Heading 2
    doc.add_paragraph(field_name, style="Heading 2")
    # 描述文字（灰色斜体提示）
    p = doc.add_paragraph()
    run = p.add_run(f"[{field_desc}]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()  # 空行

# 保存
output_path = "/Users/wulei/.openclaw/workspace/产品版本发布说明模板.docx"
doc.save(output_path)
print(f"✅ 文档已生成: {output_path}")
