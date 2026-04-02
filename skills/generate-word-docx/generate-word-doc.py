import os
import sys
import subprocess
import importlib

# --- 🛠️ 第一部分：自动检测与安装依赖 ---
def install_package(package_name):
    """
    自动调用 pip 安装指定的库
    """
    print(f"⚠️ 检测到缺少库 '{package_name}'，正在自动安装...")
    try:
        # 使用 sys.executable 确保安装到当前运行的 Python 环境中
        subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])
        print(f"✅ 库 '{package_name}' 安装成功！")
    except subprocess.CalledProcessError:
        print(f"❌ 安装失败，请手动运行: pip install {package_name}")
        sys.exit(1)

# 尝试导入 docx，如果失败则安装
try:
    import docx
except ImportError:
    # 注意：pip 安装名是 'python-docx'，但导入名是 'docx'
    install_package("python-docx")
    import docx  # 安装后再次导入

# 导入依赖库中的具体模块（必须在确保 docx 安装后进行）
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 📝 第二部分：核心功能代码 ---

def create_word_file(filename, content_data):
    """
    创建 Word 文档的主函数
    """
    # 1. 文件名处理
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    print(f"📄 正在创建文档: {filename} ...")

    # 2. 创建文档对象
    doc = Document()

    # 3. 循环写入内容
    for item in content_data:
        text = item.get('text', '')
        style = item.get('style', 'Normal')
        is_bold = item.get('bold', False)
        hex_color = item.get('color', None)
        alignment = item.get('alignment', 'LEFT')

        # 添加段落 (增加容错，防止样式不存在报错)
        try:
            paragraph = doc.add_paragraph(text, style=style)
        except KeyError:
            paragraph = doc.add_paragraph(text, style='Normal')

        # 设置对齐
        if alignment == 'CENTER':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'RIGHT':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 设置字体格式
        run = paragraph.runs[0]
        if is_bold:
            run.bold = True
        
        if hex_color:
            # 清洗颜色代码，移除可能存在的 '#'
            clean_hex = hex_color.lstrip('#')
            if len(clean_hex) == 6:
                try:
                    r = int(clean_hex[:2], 16)
                    g = int(clean_hex[2:4], 16)
                    b = int(clean_hex[4:], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except ValueError:
                    pass

    # 4. 保存文件
    try:
        doc.save(filename)
        abs_path = os.path.abspath(filename)
        print(f"✅ 成功！文件已保存至: {abs_path}")
        return abs_path
    except PermissionError:
        print(f"❌ 错误: 无法保存。请关闭已打开的 '{filename}' 文档后重试。")
    except Exception as e:
        print(f"❌ 未知错误: {e}")

# --- 🚀 第三部分：执行入口 ---
if __name__ == "__main__":
    # 定义测试数据
    file_name = "自动安装测试.docx"
    contents = [
        {"text": "环境自动配置测试", "style": "Heading 1", "alignment": "CENTER"},
        {"text": "如果你看到这段文字，说明 python-docx 库已经自动安装成功了！", "style": "Normal"},
        {"text": "红色加粗文字", "style": "Normal", "color": "#FF0000", "bold": True}
    ]
    
    # 运行函数
    create_word_file(file_name, contents)
